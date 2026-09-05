"""T-339 — independent structural verification of the exported Word document.

Validates the editor's export with tooling that shares no code with the editor:
python-docx must open the file (proves it is a well-formed package a Word
processor accepts), and raw OOXML asserts prove every edit family from the
task contract survived export.

Usage: python3 scripts/verify-t339-word-export.py [path-to-export]
"""

import json
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import docx

DOCX = Path(sys.argv[1] if len(sys.argv) > 1 else "artifacts/t-339-word/r1/export-t339-word.docx")
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
PKG = "{http://schemas.openxmlformats.org/package/2006/relationships}"

# 1. The package must be a valid zip and python-docx must open it.
with zipfile.ZipFile(DOCX) as archive:
    assert archive.testzip() is None, "zip integrity check failed"
    document = ET.fromstring(archive.read("word/document.xml"))
    relationships = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
    numbering = ET.fromstring(archive.read("word/numbering.xml"))
    content_types = archive.read("[Content_Types].xml").decode()
    media = [n for n in archive.namelist() if n.startswith("word/media/") and not n.endswith("/")]

opened = docx.Document(str(DOCX))  # raises if the file is not a valid Word document

body_text = "".join(node.text or "" for node in document.iter(f"{W}t"))
paragraphs = document.findall(f"{W}body/{W}p")
styles = [node.attrib.get(f"{W}val") for node in document.iter(f"{W}pStyle")]
tables = document.findall(f"{W}body/{W}tbl")
rel_targets = {
    rel.attrib["Id"]: rel.attrib["Target"] for rel in relationships.findall(f"{PKG}Relationship")
}


def paragraph_with(text):
    for p in paragraphs:
        if text in "".join(node.text or "" for node in p.iter(f"{W}t")):
            return p
    raise AssertionError(f"paragraph not found: {text}")


report = {}

# 2. Text add / edit / delete
report["editedBody"] = "T339-BODY-EDITED line-A" in body_text
report["originalBodyGone"] = "T339-BODY-ORIGINAL" not in body_text
report["throwawayDeleted"] = "T339-DELETE-ME" not in body_text
# multi-line edit must be explicit w:br runs, never raw newlines in w:t
edited = paragraph_with("T339-BODY-EDITED line-A")
report["multilineAsBr"] = len(edited.findall(f"{W}r/{W}br")) == 1 and not any(
    "\n" in (node.text or "") for node in edited.iter(f"{W}t")
)

# 3. Copy / paste: the CJK sentinel paragraph must exist twice
report["cjkPastedTwice"] = body_text.count("T339-CJK") == 2

# 4. Styles, font, alignment, bold
report["heading3Applied"] = "Heading3" in styles
ref = paragraph_with("Reference material follows.")
assert ref.find(f"{W}pPr/{W}pStyle").attrib[f"{W}val"] == "Heading3"
cjk = paragraph_with("T339-CJK")
fonts = cjk.find(f"{W}r/{W}rPr/{W}rFonts")
size = cjk.find(f"{W}r/{W}rPr/{W}sz")
report["fontFamily"] = fonts is not None and fonts.attrib.get(f"{W}ascii") == "Georgia"
report["fontSizeHalfPoints"] = size is not None and size.attrib.get(f"{W}val") == "28"
title = paragraph_with("Word fidelity fixture")
report["titleCentered"] = title.find(f"{W}pPr/{W}jc").attrib[f"{W}val"] == "center"
report["bodyJustified"] = edited.find(f"{W}pPr/{W}jc").attrib[f"{W}val"] == "both"
closing = paragraph_with("T339-DOC-END")
report["boldToggled"] = closing.find(f"{W}r/{W}rPr/{W}b") is not None

# 5. Lists: toggled paragraphs carry numPr AND numbering.xml defines the ids
cjk_copy = [
    p
    for p in paragraphs
    if "T339-CJK" in "".join(node.text or "" for node in p.iter(f"{W}t"))
][-1]
bullet_num_id = cjk_copy.find(f"{W}pPr/{W}numPr/{W}numId").attrib[f"{W}val"]
numbered = paragraph_with("T339-NUMBERED-ITEM")
number_num_id = numbered.find(f"{W}pPr/{W}numPr/{W}numId").attrib[f"{W}val"]
defined_nums = {}
abstract_formats = {}
for abstract in numbering.findall(f"{W}abstractNum"):
    fmt = abstract.find(f"{W}lvl/{W}numFmt")
    abstract_formats[abstract.attrib[f"{W}abstractNumId"]] = fmt.attrib[f"{W}val"]
for num in numbering.findall(f"{W}num"):
    abstract_id = num.find(f"{W}abstractNumId").attrib[f"{W}val"]
    defined_nums[num.attrib[f"{W}numId"]] = abstract_formats.get(abstract_id)
report["bulletNumberingDefined"] = defined_nums.get(bullet_num_id) == "bullet"
report["decimalNumberingDefined"] = defined_nums.get(number_num_id) == "decimal"
report["numberingContentTypeRegistered"] = "/word/numbering.xml" in content_types
report["numberingRelRegistered"] = "numbering.xml" in rel_targets.values()

# 6. Tables: edited cell + appended 3x3 table
def table_cells(table):
    return [
        [
            "".join(node.text or "" for node in cell.iter(f"{W}t"))
            for cell in row.findall(f"{W}tc")
        ]
        for row in table.findall(f"{W}tr")
    ]

report["tableCount"] = len(tables)
report["cellEdited"] = table_cells(tables[0])[2][2] == "Ready T339"
report["appendedTable3x3"] = len(table_cells(tables[1])) == 3 and len(table_cells(tables[1])[0]) == 3

# 7. Hyperlinks: original + added, both resolvable, added one visibly styled
hyperlinks = list(document.iter(f"{W}hyperlink"))
targets = [rel_targets.get(h.attrib.get(f"{R}id")) for h in hyperlinks]
report["hyperlinkTargets"] = targets
assert "https://github.com/lobehub/lobehub" in targets
assert "https://lobehub.com" in targets
added = hyperlinks[[i for i, t in enumerate(targets) if t == "https://lobehub.com"][0]]
report["addedLinkStyled"] = (
    added.find(f"{W}r/{W}rPr/{W}color") is not None
    and added.find(f"{W}r/{W}rPr/{W}u") is not None
)

# 8. Images: original replaced in place + newly inserted picture with drawing XML
report["mediaParts"] = media
report["drawingCount"] = len(list(document.iter(f"{W}drawing")))
assert report["drawingCount"] == 2, "expected original + inserted drawing"
assert any("lobehub-image-" in name for name in media), "inserted media part missing"
inline_ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
extents = [e.attrib for e in document.iter(f"{inline_ns}extent")]
report["insertedImageExtent"] = extents
# the replaced original image bytes must differ from the fixture's blue PNG
fixture = Path("docs/development/t-339-word/fixtures/t339-word-fixture.docx")
with zipfile.ZipFile(DOCX) as archive, zipfile.ZipFile(fixture) as source:
    report["originalImageReplaced"] = archive.read("word/media/image1.png") != source.read(
        "word/media/image1.png"
    )

# 9. CJK fidelity
report["cjkIntact"] = "中文保真段落：排版与导出必须无损" in body_text

# python-docx object-level sanity: it parsed paragraphs and tables
report["pythonDocxParagraphs"] = len(opened.paragraphs)
report["pythonDocxTables"] = len(opened.tables)

failures = [key for key, value in report.items() if value is False]
print(json.dumps(report, ensure_ascii=False, indent=2))
if failures:
    print("FAILED:", failures, file=sys.stderr)
    sys.exit(1)
print("ALL CHECKS PASSED")
