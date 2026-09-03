from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw


OUTPUT = Path('artifacts/t-328-word/word-roundtrip-exported.docx')


def hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    link = OxmlElement('w:hyperlink')
    link.set(qn('r:id'), relationship_id)
    run = OxmlElement('w:r')
    properties = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    properties.extend([color, underline])
    value = OxmlElement('w:t')
    value.text = text
    run.extend([properties, value])
    link.append(run)
    paragraph._p.append(link)


def page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = ' PAGE '
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instruction, end])


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(2.5)

    normal = document.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    for name, size, color in [('Title', 26, '17324D'), ('Heading 1', 18, '17324D'), ('Heading 2', 14, '1C6E8C')]:
        style = document.styles[name]
        style.font.name = 'Arial'
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    document.core_properties.title = 'LH-OFFICE-V1-DOCX'
    document.add_paragraph('Office round-trip fixture', style='Title')
    document.add_heading('LH-OFFICE-V1-DOC-H1', level=1)
    first = document.add_paragraph()
    first.add_run('This body paragraph keeps ').bold = False
    first.add_run('bold').bold = True
    first.add_run(', ')
    first.add_run('italic').italic = True
    first.add_run(', and ')
    first.add_run('underlined').underline = True
    first.add_run(' runs after export.')
    pasted = document.add_paragraph('Pasted body paragraph — T-328-PASTE', style='Body Text')
    pasted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_paragraph('中文保真检查 ✅ — edited')

    document.add_heading('Checklist', level=1)
    for item in ['Import the source', 'Edit content and formatting', 'Save and reopen', 'Export the DOCX']:
        document.add_paragraph(item, style='List Number')
    for item in ['No content loss', 'No layout drift']:
        document.add_paragraph(item, style='List Bullet')

    table = document.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for cell, value in zip(table.rows[0].cells, ['Item', 'Owner', 'Status']):
        cell.text = value
        cell.paragraphs[0].runs[0].bold = True
    for row, values in zip(table.rows[1:], [('Draft', 'Ari', 'Ready'), ('Export', 'Bo', 'Ready')]):
        for cell, value in zip(row.cells, values):
            cell.text = value

    image = Image.new('RGBA', (640, 360), (245, 248, 250, 255))
    draw = ImageDraw.Draw(image)
    draw.rectangle((20, 20, 150, 150), fill=(220, 55, 55, 255))
    draw.ellipse((450, 190, 610, 350), fill=(40, 95, 190, 255))
    draw.text((205, 165), 'LH-OFFICE-V1-IMAGE', fill=(23, 50, 77, 255))
    image_buffer = BytesIO()
    image.save(image_buffer, format='PNG')
    image_buffer.seek(0)
    picture = document.add_picture(image_buffer, width=Cm(12))
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    properties = picture._inline.docPr
    properties.set('descr', 'LH-OFFICE-V1-IMAGE-ALT')
    caption = document.add_paragraph('Figure 1 — Synthetic fidelity image')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.style = 'Caption'

    link_paragraph = document.add_paragraph('Source: ')
    hyperlink(link_paragraph, 'LobeHub repository', 'https://github.com/lobehub/lobehub')
    link_paragraph.add_run(' (target preserved).')

    document.add_page_break()
    document.add_heading('Appendix', level=1)
    document.add_paragraph('LH-OFFICE-V1-DOC-END')
    document.add_paragraph('Reopen/export invariant: BODY + HIERARCHY + TABLE + IMAGE + LINK.')

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('LH-OFFICE-V1-DOC-FOOTER · Page ')
    page_number(footer)
    document.save(OUTPUT)
    print(f'CREATED {OUTPUT} bytes={OUTPUT.stat().st_size}')


if __name__ == '__main__':
    main()
