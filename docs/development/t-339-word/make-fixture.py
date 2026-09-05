"""T-339 — builds the multi-paragraph Word fixture the editor battery runs against.

Deliberately authored with python-docx (not the editor under test) so the input
document is produced by independent tooling. Covers every content family the
task contract names: headings, body paragraphs (incl. CJK), bullet + numbered
lists, a table, an inline image, and an external hyperlink.

Usage: python3 docs/development/t-339-word/make-fixture.py
Output: docs/development/t-339-word/fixtures/t339-word-fixture.docx
"""

import io
import struct
import zlib
from pathlib import Path

import docx
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches

OUT = Path(__file__).parent / "fixtures" / "t339-word-fixture.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)


def make_png(width: int, height: int, rgb: tuple) -> bytes:
    """Hand-rolled solid-color PNG so the fixture needs no imaging library."""
    raw = b"".join(b"\x00" + bytes(rgb) * width for _ in range(height))

    def chunk(kind: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + kind
            + payload
            + struct.pack(">I", zlib.crc32(kind + payload) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_props)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


document = docx.Document()
document.add_paragraph("T-339 Word fidelity fixture", style="Title")
document.add_paragraph("Section overview", style="Heading 1")
document.add_paragraph(
    "This body paragraph exists to be edited, copied, and re-styled by the "
    "LobeHub Word editor battery. Sentinel: T339-BODY-ORIGINAL."
)
document.add_paragraph("中文保真段落：排版与导出必须无损。Sentinel: T339-CJK。")
document.add_paragraph("Formatting targets", style="Heading 2")
document.add_paragraph("Bullet alpha", style="List Bullet")
document.add_paragraph("Bullet beta", style="List Bullet")
document.add_paragraph("Numbered one", style="List Number")
document.add_paragraph("Numbered two", style="List Number")

table = document.add_table(rows=3, cols=3)
table.style = "Table Grid"
for column, heading in enumerate(["Item", "Owner", "Status"]):
    table.cell(0, column).text = heading
table.cell(1, 0).text = "Import"
table.cell(1, 1).text = "Ada"
table.cell(1, 2).text = "Done"
table.cell(2, 0).text = "Export"
table.cell(2, 1).text = "Bo"
table.cell(2, 2).text = "Pending"

document.add_paragraph("Reference material follows.")
link_paragraph = document.add_paragraph()
add_hyperlink(link_paragraph, "https://github.com/lobehub/lobehub", "LobeHub repository")

document.add_picture(io.BytesIO(make_png(96, 64, (52, 120, 246))), width=Inches(2))
document.add_paragraph("Closing sentinel: T339-DOC-END.")

document.save(OUT)
print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")
